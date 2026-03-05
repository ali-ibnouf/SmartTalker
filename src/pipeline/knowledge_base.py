"""Knowledge Base engine with RAG (Retrieval-Augmented Generation).

Uses DashScope embeddings (text-embedding-v3) + ChromaDB for document storage and retrieval.
Supports PDF, TXT, CSV, DOCX, JSON document ingestion.
"""

from __future__ import annotations

import csv
import hashlib
import io
import json
import time
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

import httpx

from src.config import Settings
from src.utils.exceptions import KnowledgeBaseError
from src.utils.logger import setup_logger, log_with_latency

logger = setup_logger("pipeline.knowledge_base")


# ── Dataclasses ──────────────────────────────────────────────────────────────


@dataclass
class KBDocument:
    """Metadata for an ingested document."""

    doc_id: str
    filename: str
    doc_type: str
    chunk_count: int = 0
    created_at: float = 0.0
    file_hash: str = ""
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class KBSearchResult:
    """Result of a knowledge base vector search."""

    chunks: list[dict[str, Any]] = field(default_factory=list)
    query: str = ""
    top_similarity: float = 0.0
    latency_ms: int = 0


@dataclass
class KBQueryResult:
    """Result of a RAG query (search + context assembly).

    Attributes:
        context: Assembled context string for LLM injection.
        source_chunks: Source chunks that formed the context.
        confidence: Confidence that KB can answer (based on similarity).
        has_answer: Whether KB has relevant context above threshold.
        latency_ms: Total query processing time.
    """

    context: str = ""
    source_chunks: list[dict[str, Any]] = field(default_factory=list)
    confidence: float = 0.0
    has_answer: bool = False
    latency_ms: int = 0


# ── Engine ───────────────────────────────────────────────────────────────────


class KnowledgeBaseEngine:
    """Knowledge Base engine with document ingestion and RAG retrieval.

    Follows the standard SmartTalker engine pattern:
    - ``__init__(config)`` / ``is_loaded`` / ``load()`` / ``unload()``
    - Uses DashScope OpenAI-compatible API for embeddings (text-embedding-v3)
    - Uses ChromaDB for vector storage
    """

    def __init__(self, config: Settings) -> None:
        self._config = config
        self._loaded = False
        self._collection: Any = None
        self._chroma_client: Any = None
        self._http_client: Optional[httpx.AsyncClient] = None
        self._docs: dict[str, KBDocument] = {}

        logger.info("KnowledgeBaseEngine initialized")

    @property
    def is_loaded(self) -> bool:
        return self._loaded

    # ── Lifecycle ─────────────────────────────────────────────────────

    def load(self) -> None:
        """Load ChromaDB persistent client and collection."""
        try:
            import chromadb
        except ImportError:
            raise KnowledgeBaseError(
                message="chromadb not installed",
                detail="Install with: pip install chromadb>=0.5.0",
            )

        try:
            persist_dir = str(self._config.kb_storage_dir / "chroma")
            Path(persist_dir).mkdir(parents=True, exist_ok=True)

            self._chroma_client = chromadb.PersistentClient(path=persist_dir)
            self._collection = self._chroma_client.get_or_create_collection(
                name="smarttalker_kb",
                metadata={"hnsw:space": "cosine"},
            )

            # Rebuild _docs index from ChromaDB metadata
            self._rebuild_docs_index()

            self._loaded = True
            logger.info(
                "KnowledgeBaseEngine loaded",
                extra={"persist_dir": persist_dir, "doc_count": len(self._docs)},
            )
        except Exception as exc:
            raise KnowledgeBaseError(
                message="Failed to initialize Knowledge Base",
                detail=str(exc),
                original_exception=exc,
            ) from exc

    def unload(self) -> None:
        """Free resources."""
        if self._http_client is not None:
            # Schedule close if possible; sync fallback
            try:
                import asyncio

                loop = asyncio.get_running_loop()
                loop.create_task(self._http_client.aclose())
            except RuntimeError:
                pass
            self._http_client = None

        self._collection = None
        self._chroma_client = None
        self._docs.clear()
        self._loaded = False
        logger.info("KnowledgeBaseEngine unloaded")

    # ── Document Ingestion ────────────────────────────────────────────

    async def ingest_document(
        self,
        file_path: str,
        doc_type: str,
        metadata: Optional[dict[str, Any]] = None,
    ) -> KBDocument:
        """Ingest a document: parse, chunk, embed, store in ChromaDB."""
        if not self._loaded:
            raise KnowledgeBaseError(message="KB engine not loaded")

        start = time.perf_counter()
        path = Path(file_path)
        if not path.exists():
            raise KnowledgeBaseError(message=f"File not found: {file_path}")

        # Parse
        parsers = {
            "pdf": self._parse_pdf,
            "txt": self._parse_txt,
            "csv": self._parse_csv,
            "docx": self._parse_docx,
            "json": self._parse_json,
        }
        parser = parsers.get(doc_type)
        if not parser:
            raise KnowledgeBaseError(
                message=f"Unsupported document type: {doc_type}",
                detail=f"Supported types: {', '.join(parsers.keys())}",
            )

        raw_text = parser(file_path)
        if not raw_text.strip():
            raise KnowledgeBaseError(message="Document produced no text content")

        # Chunk
        chunks = self._chunk_text(raw_text)

        # Compute file hash
        file_hash = hashlib.sha256(path.read_bytes()).hexdigest()[:16]

        # Embed
        embeddings = await self._get_embeddings_batch(chunks)

        # Store in ChromaDB
        doc_id = f"doc_{uuid.uuid4().hex[:12]}"
        ids = [f"{doc_id}_chunk_{i}" for i in range(len(chunks))]
        metadatas = [
            {
                "doc_id": doc_id,
                "chunk_index": i,
                "filename": path.name,
                "doc_type": doc_type,
                **(metadata or {}),
            }
            for i in range(len(chunks))
        ]

        self._collection.add(
            ids=ids,
            embeddings=embeddings,
            documents=chunks,
            metadatas=metadatas,
        )

        elapsed_ms = int((time.perf_counter() - start) * 1000)
        doc = KBDocument(
            doc_id=doc_id,
            filename=path.name,
            doc_type=doc_type,
            chunk_count=len(chunks),
            created_at=time.time(),
            file_hash=file_hash,
            metadata=metadata or {},
        )
        self._docs[doc_id] = doc

        log_with_latency(
            logger,
            "Document ingested",
            elapsed_ms,
            extra={"doc_id": doc_id, "chunks": len(chunks), "filename": path.name},
        )
        return doc

    async def ingest_text(
        self,
        text: str,
        source_name: str,
        metadata: Optional[dict[str, Any]] = None,
    ) -> KBDocument:
        """Ingest raw text directly (used for training Q&A pairs)."""
        if not self._loaded:
            raise KnowledgeBaseError(message="KB engine not loaded")

        if not text.strip():
            raise KnowledgeBaseError(message="Text is empty")

        start = time.perf_counter()
        chunks = self._chunk_text(text)
        embeddings = await self._get_embeddings_batch(chunks)

        doc_id = f"txt_{uuid.uuid4().hex[:12]}"
        ids = [f"{doc_id}_chunk_{i}" for i in range(len(chunks))]
        metadatas = [
            {
                "doc_id": doc_id,
                "chunk_index": i,
                "filename": source_name,
                "doc_type": "text",
                **(metadata or {}),
            }
            for i in range(len(chunks))
        ]

        self._collection.add(
            ids=ids,
            embeddings=embeddings,
            documents=chunks,
            metadatas=metadatas,
        )

        elapsed_ms = int((time.perf_counter() - start) * 1000)
        doc = KBDocument(
            doc_id=doc_id,
            filename=source_name,
            doc_type="text",
            chunk_count=len(chunks),
            created_at=time.time(),
            file_hash=hashlib.sha256(text.encode()).hexdigest()[:16],
            metadata=metadata or {},
        )
        self._docs[doc_id] = doc

        log_with_latency(
            logger,
            "Text ingested",
            elapsed_ms,
            extra={"doc_id": doc_id, "chunks": len(chunks), "source": source_name},
        )
        return doc

    def list_documents(self) -> list[KBDocument]:
        """List all ingested documents."""
        return list(self._docs.values())

    def delete_document(self, doc_id: str) -> bool:
        """Delete a document and its chunks from ChromaDB."""
        if not self._loaded:
            raise KnowledgeBaseError(message="KB engine not loaded")

        if doc_id not in self._docs:
            return False

        # Remove chunks from ChromaDB
        self._collection.delete(where={"doc_id": doc_id})
        del self._docs[doc_id]

        logger.info("Document deleted", extra={"doc_id": doc_id})
        return True

    # ── Search / RAG ──────────────────────────────────────────────────

    async def search(
        self,
        query: str,
        top_k: Optional[int] = None,
    ) -> KBSearchResult:
        """Vector search for relevant chunks."""
        if not self._loaded:
            raise KnowledgeBaseError(message="KB engine not loaded")

        start = time.perf_counter()
        k = top_k or self._config.kb_top_k

        # Check if collection has any documents
        if self._collection.count() == 0:
            return KBSearchResult(query=query, latency_ms=0)

        query_embedding = await self._get_embedding(query)
        results = self._collection.query(
            query_embeddings=[query_embedding],
            n_results=min(k, self._collection.count()),
            include=["documents", "metadatas", "distances"],
        )

        chunks = []
        top_similarity = 0.0

        if results["documents"] and results["documents"][0]:
            for i, doc_text in enumerate(results["documents"][0]):
                # ChromaDB cosine distance: 0 = identical, 2 = opposite
                # Convert to similarity: 1 - (distance / 2)
                distance = results["distances"][0][i]
                similarity = 1.0 - (distance / 2.0)

                meta = results["metadatas"][0][i] if results["metadatas"] else {}
                chunks.append({
                    "text": doc_text,
                    "similarity": similarity,
                    "filename": meta.get("filename", ""),
                    "doc_id": meta.get("doc_id", ""),
                    "chunk_index": meta.get("chunk_index", 0),
                })

                if similarity > top_similarity:
                    top_similarity = similarity

        elapsed_ms = int((time.perf_counter() - start) * 1000)
        return KBSearchResult(
            chunks=chunks,
            query=query,
            top_similarity=top_similarity,
            latency_ms=elapsed_ms,
        )

    async def query(self, user_text: str) -> KBQueryResult:
        """Full RAG query: embed, search, assemble context, score confidence."""
        if not self._loaded:
            return KBQueryResult()

        start = time.perf_counter()
        search_result = await self.search(user_text)

        if not search_result.chunks:
            elapsed_ms = int((time.perf_counter() - start) * 1000)
            return KBQueryResult(latency_ms=elapsed_ms)

        # Filter chunks above threshold
        threshold = self._config.kb_confidence_threshold
        relevant = [c for c in search_result.chunks if c["similarity"] >= threshold]

        if not relevant:
            elapsed_ms = int((time.perf_counter() - start) * 1000)
            return KBQueryResult(
                confidence=search_result.top_similarity,
                latency_ms=elapsed_ms,
            )

        # Assemble context string
        context_parts = []
        for chunk in relevant:
            source = chunk.get("filename", "unknown")
            context_parts.append(f"[Source: {source}]\n{chunk['text']}")

        context = "\n\n".join(context_parts)
        confidence = relevant[0]["similarity"]

        elapsed_ms = int((time.perf_counter() - start) * 1000)
        return KBQueryResult(
            context=context,
            source_chunks=relevant,
            confidence=confidence,
            has_answer=True,
            latency_ms=elapsed_ms,
        )

    # ── Embedding ─────────────────────────────────────────────────────

    async def _get_embedding(self, text: str) -> list[float]:
        """Get embedding from DashScope OpenAI-compatible API (POST /embeddings)."""
        if self._http_client is None or self._http_client.is_closed:
            headers = {"Content-Type": "application/json"}
            if self._config.kb_embedding_api_key:
                headers["Authorization"] = f"Bearer {self._config.kb_embedding_api_key}"
            self._http_client = httpx.AsyncClient(
                base_url=self._config.kb_embedding_api_url,
                timeout=httpx.Timeout(30.0, connect=10.0),
                headers=headers,
            )

        try:
            response = await self._http_client.post(
                "/embeddings",
                json={
                    "model": self._config.kb_embedding_model,
                    "input": text,
                },
            )
            response.raise_for_status()
            data = response.json()
            return data["data"][0]["embedding"]
        except httpx.ConnectError as exc:
            raise KnowledgeBaseError(
                message="Cannot connect to embedding API",
                detail=f"Check that DashScope API is reachable at: {self._config.kb_embedding_api_url}",
                original_exception=exc,
            ) from exc
        except Exception as exc:
            raise KnowledgeBaseError(
                message="Embedding generation failed",
                detail=str(exc),
                original_exception=exc,
            ) from exc

    async def _get_embeddings_batch(self, texts: list[str]) -> list[list[float]]:
        """Batch embedding — DashScope supports batch input."""
        if not texts:
            return []

        if self._http_client is None or self._http_client.is_closed:
            headers = {"Content-Type": "application/json"}
            if self._config.kb_embedding_api_key:
                headers["Authorization"] = f"Bearer {self._config.kb_embedding_api_key}"
            self._http_client = httpx.AsyncClient(
                base_url=self._config.kb_embedding_api_url,
                timeout=httpx.Timeout(60.0, connect=10.0),
                headers=headers,
            )

        try:
            response = await self._http_client.post(
                "/embeddings",
                json={
                    "model": self._config.kb_embedding_model,
                    "input": texts,
                },
            )
            response.raise_for_status()
            data = response.json()
            # Sort by index to ensure order matches input
            sorted_data = sorted(data["data"], key=lambda x: x["index"])
            return [item["embedding"] for item in sorted_data]
        except Exception as exc:
            # Fallback to sequential calls
            logger.warning(f"Batch embedding failed, falling back to sequential: {exc}")
            embeddings = []
            for text in texts:
                emb = await self._get_embedding(text)
                embeddings.append(emb)
            return embeddings

    # ── Document Parsing ──────────────────────────────────────────────

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """Extract text from a PDF file."""
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise KnowledgeBaseError(
                message="PyPDF2 not installed",
                detail="Install with: pip install pypdf2",
            )

        reader = PdfReader(file_path)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """Extract text from a DOCX file."""
        try:
            from docx import Document
        except ImportError:
            raise KnowledgeBaseError(
                message="python-docx not installed",
                detail="Install with: pip install python-docx",
            )

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    @staticmethod
    def _parse_csv(file_path: str) -> str:
        """Extract text from a CSV file (rows joined as text)."""
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            rows = []
            for row in reader:
                rows.append(" | ".join(row))
        return "\n".join(rows)

    @staticmethod
    def _parse_txt(file_path: str) -> str:
        """Read plain text file."""
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    @staticmethod
    def _parse_json(file_path: str) -> str:
        """Extract text from a JSON file (flattened to readable text)."""
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        def flatten(obj: Any, prefix: str = "") -> list[str]:
            lines = []
            if isinstance(obj, dict):
                for k, v in obj.items():
                    lines.extend(flatten(v, f"{prefix}{k}: "))
            elif isinstance(obj, list):
                for i, item in enumerate(obj):
                    lines.extend(flatten(item, f"{prefix}[{i}] "))
            else:
                lines.append(f"{prefix}{obj}")
            return lines

        return "\n".join(flatten(data))

    # ── Chunking ──────────────────────────────────────────────────────

    def _chunk_text(self, text: str) -> list[str]:
        """Split text into chunks with overlap."""
        chunk_size = self._config.kb_chunk_size
        overlap = self._config.kb_chunk_overlap

        if len(text) <= chunk_size:
            return [text.strip()] if text.strip() else []

        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            start += chunk_size - overlap

        return chunks

    # ── Internal ──────────────────────────────────────────────────────

    def _rebuild_docs_index(self) -> None:
        """Rebuild the in-memory document index from ChromaDB metadata."""
        if self._collection is None or self._collection.count() == 0:
            return

        all_data = self._collection.get(include=["metadatas"])
        seen_docs: dict[str, dict[str, Any]] = {}

        for meta in all_data["metadatas"]:
            doc_id = meta.get("doc_id", "")
            if doc_id and doc_id not in seen_docs:
                seen_docs[doc_id] = {
                    "filename": meta.get("filename", ""),
                    "doc_type": meta.get("doc_type", "unknown"),
                    "chunk_count": 0,
                }
            if doc_id:
                seen_docs[doc_id]["chunk_count"] += 1

        for doc_id, info in seen_docs.items():
            self._docs[doc_id] = KBDocument(
                doc_id=doc_id,
                filename=info["filename"],
                doc_type=info["doc_type"],
                chunk_count=info["chunk_count"],
                created_at=0.0,
            )
